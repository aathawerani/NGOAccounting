import calendar
from datetime import date
from io import BytesIO
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy import or_, and_
from sqlalchemy.orm import Session, joinedload

from audit import log_audit
from database import get_db
from models.models import AccountType, LedgerEntry, RentReceipt, Tenant, Trust
from pdf_utils import NGODoc, hijri_str

router = APIRouter(prefix="/api/rent", tags=["rent"])


# ── Pydantic schemas ──────────────────────────────────────────────────────────

class RentReceiptBody(BaseModel):
    trust_id: int
    tenant_id: int
    date: date
    from_month: int     # 1–12
    from_year: int
    to_month: int       # 1–12
    to_year: int
    rent_arrears: float = 0.0
    water_arrears: float = 0.0
    debit_account_code: str = "CASH"   # which cash/bank account to debit


# ── Helpers ───────────────────────────────────────────────────────────────────

def _fmt(d: date) -> str:
    """WPF particulars date format: d/m/yyyy (no zero-padding)."""
    return f"{d.day}/{d.month}/{d.year}"


def _num_months(from_m: int, from_y: int, to_m: int, to_y: int) -> int:
    """Inclusive month count: Jan→Mar = 3."""
    return (to_y - from_y) * 12 + (to_m - from_m) + 1


def _next_serial(trust_id: int, db: Session) -> str:
    """3-digit zero-padded serial, increments from highest existing per trust."""
    last = (
        db.query(RentReceipt)
        .filter(RentReceipt.trust_id == trust_id)
        .order_by(RentReceipt.id.desc())
        .first()
    )
    if not last or not last.serial_no:
        return "001"
    try:
        num = int(last.serial_no)
        nxt = (num % 999) + 1
        return f"{nxt:03d}"
    except ValueError:
        return "001"


def _serialize(r: RentReceipt) -> dict:
    return {
        "id": r.id,
        "trust_id": r.trust_id,
        "tenant_id": r.tenant_id,
        "serial_no": r.serial_no,
        "date": r.date.isoformat() if r.date else None,
        "plot_code": r.plot_code,
        "space_type": r.space_type,
        "space_number": r.space_number,
        "monthly_rent": r.monthly_rent,
        "water_charge": r.water_charge,
        "tenant_name": r.tenant_name,
        "cnic": r.cnic,
        "from_date": r.from_date.isoformat() if r.from_date else None,
        "to_date": r.to_date.isoformat() if r.to_date else None,
        "rent_arrears": r.rent_arrears,
        "water_arrears": r.water_arrears,
        "total_rent": r.total_rent,
        "total_water": r.total_water,
        "total_amount": r.total_amount,
        "rent_particulars": r.rent_particulars,
        "water_particulars": r.water_particulars,
        "arrears_particulars": r.arrears_particulars,
        "water_arrears_particulars": r.water_arrears_particulars,
    }


def _rent_water_accounts(trust_code: str, plot_code: Optional[str]) -> tuple[str, str]:
    """Return (rent_account_code, water_account_code) for a trust + plot."""
    tc = (trust_code or "").upper()
    if tc == "HVHT":
        return "RGK6", "WGK6"
    if tc == "BIB":
        return "RENT", "WSC"
    # HTTT (and fallback): R{plot} / W{plot}
    if plot_code:
        return f"R{plot_code}", f"W{plot_code}"
    return "RENT", "WSC"   # safe fallback


def _create_journal_entries(
    receipt: RentReceipt,
    trust_code: str,
    debit_code: str,
    db: Session,
) -> None:
    """Create double-entry LedgerEntry rows for rent and/or water amounts."""
    rent_code, water_code = _rent_water_accounts(trust_code, receipt.plot_code)

    rent_total  = (receipt.total_rent  or 0.0) + (receipt.rent_arrears  or 0.0)
    water_total = (receipt.total_water or 0.0) + (receipt.water_arrears or 0.0)

    entries_to_add = []

    if rent_total > 0:
        key = f"rent-{receipt.id}"
        entries_to_add += [
            LedgerEntry(
                trust_id=receipt.trust_id,
                account_code=debit_code,
                date=receipt.date,
                receipt_no=receipt.serial_no,
                party_name=receipt.tenant_name,
                contra_account_code=rent_code,
                particulars=receipt.rent_particulars,
                debit=rent_total,
                credit=0.0,
                row_order=2,
                account_key=key,
            ),
            LedgerEntry(
                trust_id=receipt.trust_id,
                account_code=rent_code,
                date=receipt.date,
                receipt_no=receipt.serial_no,
                party_name=receipt.tenant_name,
                contra_account_code=debit_code,
                particulars=receipt.rent_particulars,
                debit=0.0,
                credit=rent_total,
                row_order=2,
                account_key=key,
            ),
        ]

    if water_total > 0:
        key = f"watc-{receipt.id}"
        entries_to_add += [
            LedgerEntry(
                trust_id=receipt.trust_id,
                account_code=debit_code,
                date=receipt.date,
                receipt_no=receipt.serial_no,
                party_name=receipt.tenant_name,
                contra_account_code=water_code,
                particulars=receipt.water_particulars,
                debit=water_total,
                credit=0.0,
                row_order=2,
                account_key=key,
            ),
            LedgerEntry(
                trust_id=receipt.trust_id,
                account_code=water_code,
                date=receipt.date,
                receipt_no=receipt.serial_no,
                party_name=receipt.tenant_name,
                contra_account_code=debit_code,
                particulars=receipt.water_particulars,
                debit=0.0,
                credit=water_total,
                row_order=2,
                account_key=key,
            ),
        ]

    if entries_to_add:
        db.add_all(entries_to_add)
        db.commit()


def _delete_journal_entries(receipt_id: int, db: Session) -> None:
    """Delete all LedgerEntry rows created for a specific receipt."""
    keys = [f"rent-{receipt_id}", f"watc-{receipt_id}"]
    db.query(LedgerEntry).filter(LedgerEntry.account_key.in_(keys)).delete(synchronize_session=False)
    db.commit()


def _recalc_last_paid(tenant_id: int, db: Session):
    """After a delete, update tenant.last_paid to the latest remaining receipt."""
    tenant = db.query(Tenant).filter(Tenant.id == tenant_id).first()
    if not tenant:
        return
    latest = (
        db.query(RentReceipt)
        .filter(RentReceipt.tenant_id == tenant_id)
        .order_by(RentReceipt.to_date.desc())
        .first()
    )
    if latest and latest.to_date:
        tenant.last_paid_month = latest.to_date.month
        tenant.last_paid_year = latest.to_date.year
    else:
        tenant.last_paid_month = None
        tenant.last_paid_year = None
    db.commit()


# ── Routes ────────────────────────────────────────────────────────────────────

@router.get("/cash-accounts")
def list_cash_accounts(trust_id: int, db: Session = Depends(get_db)):
    """Return ASSET account types that are cash or bank accounts for a given trust."""
    accts = (
        db.query(AccountType)
        .filter(
            AccountType.trust_id == trust_id,
            AccountType.account_type == "ASSET",
        )
        .order_by(AccountType.account_code)
        .all()
    )
    cash_like = [
        {"account_code": a.account_code, "account_name": a.account_name}
        for a in accts
        if a.account_code == "CASH"
        or a.account_code == "BOX"
        or a.account_code.upper().startswith("BANK")
    ]
    return cash_like


@router.get("")
def list_receipts(trust_id: Optional[int] = None, db: Session = Depends(get_db)):
    q = db.query(RentReceipt).options(joinedload(RentReceipt.tenant))
    if trust_id is not None:
        q = q.filter(RentReceipt.trust_id == trust_id)
    return [_serialize(r) for r in q.order_by(RentReceipt.date.desc(), RentReceipt.id.desc()).all()]


@router.get("/next-serial")
def next_serial(trust_id: int, db: Session = Depends(get_db)):
    return {"serial_no": _next_serial(trust_id, db)}


@router.post("", status_code=201)
def create_receipt(body: RentReceiptBody, db: Session = Depends(get_db)):
    # Validate trust
    trust = db.query(Trust).filter(Trust.id == body.trust_id).first()
    if not trust:
        raise HTTPException(status_code=404, detail="Trust not found")

    # Validate tenant
    tenant = db.query(Tenant).filter(Tenant.id == body.tenant_id).first()
    if not tenant:
        raise HTTPException(status_code=404, detail="Tenant not found")

    # Validate month range
    if body.from_year > body.to_year or (
        body.from_year == body.to_year and body.from_month > body.to_month
    ):
        raise HTTPException(status_code=400, detail="From date must be before or equal to To date")

    # Compute dates
    from_date = date(body.from_year, body.from_month, 1)
    last_day = calendar.monthrange(body.to_year, body.to_month)[1]
    to_date = date(body.to_year, body.to_month, last_day)

    # Compute amounts (WPF: inclusive month count)
    n = _num_months(body.from_month, body.from_year, body.to_month, body.to_year)
    total_rent = n * tenant.monthly_rent
    total_water = n * tenant.water_charge
    total_amount = total_rent + total_water + body.rent_arrears + body.water_arrears

    # WPF-format particulars strings
    space = f"{tenant.space_type or 'SPACE'} {tenant.space_number or ''}"
    rent_part = f"{space}, RENT @{int(tenant.monthly_rent)}, {_fmt(from_date)}-{_fmt(to_date)}"
    water_part = f"{space}, WATER @{int(tenant.water_charge)}, {_fmt(from_date)}-{_fmt(to_date)}"
    arrears_part = f"{space}, RENT AREARS"
    water_arrears_part = f"{space}, WATER AREARS"

    receipt = RentReceipt(
        trust_id=body.trust_id,
        tenant_id=body.tenant_id,
        serial_no=_next_serial(body.trust_id, db),
        date=body.date,
        plot_code=tenant.plot_code,
        space_type=tenant.space_type,
        space_number=tenant.space_number,
        monthly_rent=tenant.monthly_rent,
        water_charge=tenant.water_charge,
        tenant_name=tenant.name,
        cnic=tenant.cnic,
        from_date=from_date,
        to_date=to_date,
        rent_arrears=body.rent_arrears,
        water_arrears=body.water_arrears,
        total_rent=total_rent,
        total_water=total_water,
        total_amount=total_amount,
        rent_particulars=rent_part,
        water_particulars=water_part,
        arrears_particulars=arrears_part,
        water_arrears_particulars=water_arrears_part,
    )
    db.add(receipt)

    # Update tenant last paid
    tenant.last_paid_month = body.to_month
    tenant.last_paid_year = body.to_year

    db.commit()
    db.refresh(receipt)

    # Create double-entry ledger entries (CASH/BANK DR, rent/water CR)
    _create_journal_entries(receipt, trust.code, body.debit_account_code, db)
    log_audit(db, "rent_receipts", "create", record_id=receipt.id, trust_id=receipt.trust_id,
              description=f"Receipt #{receipt.serial_no} created for {tenant.name} PKR {receipt.total_amount:,.0f}")
    db.commit()
    return _serialize(receipt)


@router.put("/{receipt_id}")
def update_receipt(receipt_id: int, body: RentReceiptBody, db: Session = Depends(get_db)):
    receipt = db.query(RentReceipt).filter(RentReceipt.id == receipt_id).first()
    if not receipt:
        raise HTTPException(status_code=404, detail="Receipt not found")

    tenant = db.query(Tenant).filter(Tenant.id == body.tenant_id).first()
    if not tenant:
        raise HTTPException(status_code=404, detail="Tenant not found")

    if body.from_year > body.to_year or (
        body.from_year == body.to_year and body.from_month > body.to_month
    ):
        raise HTTPException(status_code=400, detail="From date must be before or equal to To date")

    from_date = date(body.from_year, body.from_month, 1)
    last_day = calendar.monthrange(body.to_year, body.to_month)[1]
    to_date = date(body.to_year, body.to_month, last_day)

    n = _num_months(body.from_month, body.from_year, body.to_month, body.to_year)
    total_rent = n * tenant.monthly_rent
    total_water = n * tenant.water_charge
    total_amount = total_rent + total_water + body.rent_arrears + body.water_arrears

    space = f"{tenant.space_type or 'SPACE'} {tenant.space_number or ''}"
    receipt.date = body.date
    receipt.tenant_id = body.tenant_id
    receipt.plot_code = tenant.plot_code
    receipt.space_type = tenant.space_type
    receipt.space_number = tenant.space_number
    receipt.monthly_rent = tenant.monthly_rent
    receipt.water_charge = tenant.water_charge
    receipt.tenant_name = tenant.name
    receipt.cnic = tenant.cnic
    receipt.from_date = from_date
    receipt.to_date = to_date
    receipt.rent_arrears = body.rent_arrears
    receipt.water_arrears = body.water_arrears
    receipt.total_rent = total_rent
    receipt.total_water = total_water
    receipt.total_amount = total_amount
    receipt.rent_particulars = f"{space}, RENT @{int(tenant.monthly_rent)}, {_fmt(from_date)}-{_fmt(to_date)}"
    receipt.water_particulars = f"{space}, WATER @{int(tenant.water_charge)}, {_fmt(from_date)}-{_fmt(to_date)}"
    receipt.arrears_particulars = f"{space}, RENT AREARS"
    receipt.water_arrears_particulars = f"{space}, WATER AREARS"

    db.commit()
    _recalc_last_paid(body.tenant_id, db)
    db.refresh(receipt)

    # Refresh journal entries (delete old, create new)
    _delete_journal_entries(receipt_id, db)
    trust = db.query(Trust).filter(Trust.id == body.trust_id).first()
    if trust:
        _create_journal_entries(receipt, trust.code, body.debit_account_code, db)

    return _serialize(receipt)


@router.get("/ledger-receipts")
def ledger_receipts(trust_id: int, db: Session = Depends(get_db)):
    """
    Return rent/water receipts that were imported from WPF ledger sheets.
    These live in ledger_entries where the CONTRA account is an R or W income account
    (not WT* water-tax expense accounts or WHT withholding-tax accounts).
    """
    entries = (
        db.query(LedgerEntry)
        .filter(
            LedgerEntry.trust_id == trust_id,
            LedgerEntry.is_deleted == False,
            LedgerEntry.party_name != None,
            LedgerEntry.party_name != "",
            or_(
                LedgerEntry.contra_account_code.like("R%"),
                and_(
                    LedgerEntry.contra_account_code.like("W%"),
                    ~LedgerEntry.contra_account_code.like("WT%"),
                    LedgerEntry.contra_account_code != "WHT",
                ),
            ),
        )
        .order_by(LedgerEntry.party_name, LedgerEntry.date, LedgerEntry.id)
        .all()
    )
    return [
        {
            "id":                  e.id,
            "date":                e.date.isoformat() if e.date else None,
            "receipt_no":          e.receipt_no,
            "contra_account_code": e.contra_account_code,
            "property":            e.contra_account_code[1:] if e.contra_account_code and len(e.contra_account_code) > 1 else None,
            "entry_type":          "water" if (e.contra_account_code or "")[:1].upper() == "W" else "rent",
            "party_name":          e.party_name,
            "amount":              e.debit if e.debit else e.credit,
            "particulars":         e.particulars,
        }
        for e in entries
    ]


@router.delete("/{receipt_id}", status_code=204)
def delete_receipt(receipt_id: int, db: Session = Depends(get_db)):
    receipt = db.query(RentReceipt).filter(RentReceipt.id == receipt_id).first()
    if not receipt:
        raise HTTPException(status_code=404, detail="Receipt not found")
    tenant_id = receipt.tenant_id
    _delete_journal_entries(receipt_id, db)
    log_audit(db, "rent_receipts", "delete", record_id=receipt_id, trust_id=receipt.trust_id,
              description=f"Receipt #{receipt.serial_no} deleted for {receipt.tenant_name} PKR {receipt.total_amount:,.0f}")
    db.delete(receipt)
    db.commit()
    if tenant_id:
        _recalc_last_paid(tenant_id, db)


@router.get("/receipt/{receipt_id}/print")
def print_receipt(receipt_id: int, db: Session = Depends(get_db)):
    """Return a .docx rent receipt for download/print."""
    r = (
        db.query(RentReceipt)
        .options(joinedload(RentReceipt.trust))
        .filter(RentReceipt.id == receipt_id)
        .first()
    )
    if not r:
        raise HTTPException(status_code=404, detail="Receipt not found")

    trust_name = r.trust.name if r.trust else "Trust"
    trust_code = r.trust.code if r.trust else ""

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    def _para(text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        return p

    def _row(label, value):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r1 = p.add_run(f"{label:<30}")
        r1.font.size = Pt(10)
        r2 = p.add_run(value)
        r2.bold = True
        r2.font.size = Pt(10)

    # Header
    _para(trust_name, bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _para("RENT RECEIPT", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    doc.add_paragraph()

    # Receipt meta
    receipt_date = r.date.strftime("%-d %B %Y") if r.date else "—"
    _row("Receipt No.", f"  {r.serial_no or '—'}")
    _row("Date.", f"  {receipt_date}")
    doc.add_paragraph()

    # Tenant details
    _row("Received from.", f"  {r.tenant_name or '—'}")
    _row("CNIC.", f"  {r.cnic or '—'}")
    _row("Property.", f"  {r.space_type or ''} {r.space_number or ''}, {r.plot_code or ''}")
    doc.add_paragraph()

    # Amounts
    PKR = lambda n: f"PKR {int(n or 0):,}"

    if r.rent_particulars:
        _row("Rent.", f"  {r.rent_particulars}    {PKR(r.total_rent)}")
    if r.water_particulars and r.total_water:
        _row("Water Charges.", f"  {r.water_particulars}    {PKR(r.total_water)}")
    if r.rent_arrears:
        _row("Rent Arrears.", f"  {r.arrears_particulars or ''}    {PKR(r.rent_arrears)}")
    if r.water_arrears:
        _row("Water Arrears.", f"  {r.water_arrears_particulars or ''}    {PKR(r.water_arrears)}")

    doc.add_paragraph()
    _row("TOTAL AMOUNT.", f"  {PKR(r.total_amount)}")
    doc.add_paragraph()
    doc.add_paragraph()

    # Signature lines
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(24)
    r1 = p.add_run(f"For {trust_code or trust_name}")
    r1.font.size = Pt(10)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run("Authorised Signatory")
    r2.font.size = Pt(9)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"receipt_{r.serial_no or receipt_id}_{trust_code}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/receipt/{receipt_id}/pdf")
def receipt_pdf(receipt_id: int, db: Session = Depends(get_db)):
    """Return an A5 PDF rent receipt — opens inline in the browser."""
    r = (
        db.query(RentReceipt)
        .options(joinedload(RentReceipt.trust))
        .filter(RentReceipt.id == receipt_id)
        .first()
    )
    if not r:
        raise HTTPException(404, "Receipt not found")

    trust_name = r.trust.name if r.trust else "Trust"
    trust_code = r.trust.code if r.trust else ""
    doc_date = r.date.strftime("%d %B %Y") if r.date else "—"
    h_date = hijri_str(r.date) if r.date else ""

    doc = NGODoc(trust_name, trust_code, "RENT RECEIPT")
    doc.add_header(
        doc_number=r.serial_no or str(receipt_id),
        doc_date=doc_date,
        hijri_date=h_date,
    )
    doc.add_kv_table([
        ("Received From", r.tenant_name or "—"),
        ("CNIC", r.cnic or "—"),
        ("Property", f"{r.space_type or ''} {r.space_number or ''}, {r.plot_code or ''}".strip(", ")),
    ])

    # Build line items
    items = []
    if r.rent_particulars and (r.total_rent or 0) > 0:
        items.append(("Rent", r.rent_particulars, r.total_rent or 0.0))
    if r.water_particulars and (r.total_water or 0) > 0:
        items.append(("Water Charges", r.water_particulars, r.total_water or 0.0))
    if (r.rent_arrears or 0) > 0:
        items.append(("Rent Arrears", r.arrears_particulars or "", r.rent_arrears or 0.0))
    if (r.water_arrears or 0) > 0:
        items.append(("Water Arrears", r.water_arrears_particulars or "", r.water_arrears or 0.0))

    if not items:
        items = [("Rent", r.rent_particulars or "", r.total_amount or 0.0)]

    doc.add_line_items(items, total=r.total_amount or 0.0)
    doc.add_signature()
    doc.add_footer_note(f"Generated by NGO Accounting System · {trust_name}")

    buf = doc.build()
    fname = f"receipt_{r.serial_no or receipt_id}.pdf"
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'inline; filename="{fname}"'},
    )
